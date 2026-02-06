"""Template analyzer strategy.

Analyzes Word templates to detect dynamic variables using regex patterns
or LLM-based analysis with Few-Shot Chain-of-Thought prompting.
"""

import asyncio
import json
import logging
import re
from pathlib import Path
from typing import Any

from app.interfaces.template import BaseTemplateAnalyzer
from app.strategies.template_engine.models import DetectedVariable

logger = logging.getLogger(__name__)


# =============================================================================
# Few-Shot Chain-of-Thought Prompt for LLM Analysis
# =============================================================================

TEMPLATE_ANALYSIS_PROMPT = """
You are an expert Document Intelligence Engineer. Analyze the text segment.

DEFINITIONS:
1. Dynamic: Text that changes per client (Names, Dates, Risk Profiles, Amounts).
2. Static: Legal headers, boilerplate, firm branding.

FEW-SHOT EXAMPLES:
Input: "prepared for Mr. James Arlington on 12th March"
Output: { "is_dynamic": true, "extraction": [{ "original": "Mr. James Arlington", "var": "client_name" }, { "original": "12th March", "var": "report_date" }] }

Input: "The value of investments can go down as well as up."
Output: { "is_dynamic": false, "extraction": [] }

TASK:
Analyze the user input. Return valid JSON only.
"""


class TemplateAnalyzer(BaseTemplateAnalyzer):
    """Analyzes Word templates to detect dynamic injection points.

    Uses regex-based pattern matching for common financial document variables.
    Can be extended with LLM-based analysis for more sophisticated detection.
    """

    def __init__(
        self,
        use_llm: bool = False,
        openai_api_key: str | None = None,
        base_url: str = "https://openrouter.ai/api/v1",
        model: str = "openai/gpt-4o",
        custom_prompt: str | None = None,
    ) -> None:
        """Initialize the analyzer.

        Args:
            use_llm: Whether to use LLM for analysis (else mock/regex).
            openai_api_key: OpenAI/OpenRouter API key if use_llm=True.
            base_url: API base URL (default: OpenRouter).
            model: Model name to use for LLM analysis.
            custom_prompt: Custom prompt text to use instead of default.
        """
        self._use_llm = use_llm
        self._openai_api_key = openai_api_key
        self._base_url = base_url
        self._model = model
        self._custom_prompt = custom_prompt

        # Common patterns for UK financial documents
        self._patterns = {
            # Name patterns
            r"\b(Mr|Mrs|Ms|Dr|Prof|Sir|Dame)\s+[A-Z][a-z]+\s+[A-Z][a-z]+\b": "client_full_name",
            r"\b(Mr|Mrs|Ms|Dr|Prof)\s+[A-Z][a-z]+\b": "client_name",
            # Date patterns
            r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b": "date",
            r"\b\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b": "formatted_date",
            r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2},?\s+\d{4}\b": "short_date",
            # Monetary patterns
            r"\b£[$€]?\s?[\d,]+\.?\d{0,2}\b": "monetary_amount",
            r"\b[\d,]+\.?\d{0,2}\s?(GBP|USD|EUR)\b": "currency_amount",
            # Number patterns
            r"\b\d+%\b": "percentage",
            r"\b\d+(?:,\d{3})*(?:\.\d+)?\b": "numeric_value",
            # Address patterns
            r"\d+\s+[A-Z][a-z]+\s+(Street|Road|Lane|Drive|Avenue|Close|Way)\b": "address_line",
        }

        logger.info(
            f"TemplateAnalyzer initialized: use_llm={use_llm}, "
            f"patterns_registered={len(self._patterns)}"
        )

    async def analyze(self, file_path: str) -> list[DetectedVariable]:
        """Analyze document to detect dynamic variables.

        Uses chunked LLM analysis for faster processing - groups paragraphs
        into batches instead of processing one-by-one.

        Args:
            file_path: Path to the Word template file.

        Returns:
            List of DetectedVariable objects with paragraph indices.

        Raises:
            FileNotFoundError: If template file doesn't exist.
        """
        logger.info(f"Starting template analysis: {file_path}")

        try:
            # Validate file exists
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"Template file not found: {file_path}")

            # Load document using python-docx
            try:
                from docx import Document
            except ImportError:
                logger.error("python-docx not installed. Run: poetry add python-docx")
                raise ImportError(
                    "python-docx is required for template analysis. "
                    "Install it with: poetry add python-docx"
                )

            doc = Document(file_path)
            detected_vars: list[DetectedVariable] = []

            logger.info(f"Document loaded: {len(doc.paragraphs)} paragraphs")

            # Collect non-empty paragraphs with their indices
            paragraphs_to_analyze = [
                (idx, p.text.strip())
                for idx, p in enumerate(doc.paragraphs)
                if p.text.strip() and len(p.text.strip()) > 3  # Skip very short text
            ]

            logger.info(f"Analyzing {len(paragraphs_to_analyze)} non-empty paragraphs")

            if self._use_llm and len(paragraphs_to_analyze) > 0:
                # Use chunked LLM analysis for faster processing
                detected_vars = await self._analyze_with_llm_chunking(paragraphs_to_analyze)
            else:
                # Use regex analysis (fallback)
                for idx, text in paragraphs_to_analyze:
                    result = self._regex_analysis(text, idx)
                    detected_vars.extend(result)

            logger.info(
                f"Analysis complete: {len(detected_vars)} variables detected "
                f"from {len(doc.paragraphs)} paragraphs"
            )
            return detected_vars

        except FileNotFoundError:
            raise
        except ImportError:
            raise
        except Exception as e:
            logger.error(f"Template analysis failed: {e}", exc_info=True)
            raise RuntimeError(f"Template analysis failed: {e}") from e

    async def _analyze_with_llm_chunking(
        self, paragraphs: list[tuple[int, str]]
    ) -> list[DetectedVariable]:
        """Analyze multiple paragraphs in chunks using LLM for faster processing.

        Groups paragraphs into batches and processes them concurrently.

        Args:
            paragraphs: List of (index, text) tuples.

        Returns:
            List of all detected variables.
        """
        CHUNK_SIZE = 5  # Process 10 paragraphs per LLM call
        detected_vars: list[DetectedVariable] = []

        # Split into chunks
        chunks = [
            paragraphs[i:i + CHUNK_SIZE]
            for i in range(0, len(paragraphs), CHUNK_SIZE)
        ]

        logger.info(f"Processing {len(paragraphs)} paragraphs in {len(chunks)} chunks")

        for chunk_idx, chunk in enumerate(chunks):
            try:
                # Add timeout to prevent hanging
                async with asyncio.timeout(120):  # 120 second timeout per chunk for slower APIs
                    result = await self._call_llm_chunk(chunk, chunk_idx)
                    detected_vars.extend(result)
                    logger.debug(f"Chunk {chunk_idx + 1}/{len(chunks)} processed: {len(result)} variables")
            except asyncio.TimeoutError:
                logger.warning(f"Chunk {chunk_idx + 1} timed out, falling back to regex")
                # Fallback to regex for this chunk
                for idx, text in chunk:
                    result = self._regex_analysis(text, idx)
                    detected_vars.extend(result)
            except Exception as e:
                logger.error(f"Chunk {chunk_idx + 1} failed: {e}, falling back to regex")
                # Fallback to regex for this chunk
                for idx, text in chunk:
                    result = self._regex_analysis(text, idx)
                    detected_vars.extend(result)

        return detected_vars

    async def _call_llm_chunk(
        self, chunk: list[tuple[int, str]], chunk_idx: int
    ) -> list[DetectedVariable]:
        """Call LLM with a chunk of paragraphs for batch analysis.

        Args:
            chunk: List of (index, text) tuples.
            chunk_idx: Chunk index for logging.

        Returns:
            List of DetectedVariable objects.
        """
        if not self._openai_api_key:
            logger.warning("No OpenAI API key provided. Falling back to regex.")
            results = []
            for idx, text in chunk:
                results.extend(self._regex_analysis(text, idx))
            return results

        try:
            from openai import AsyncOpenAI
            client = AsyncOpenAI(
                api_key=self._openai_api_key,
                base_url=self._base_url,
                timeout=60.0,  # 60 second timeout
            )

            # Use batch-specific prompt for chunking
            # The custom prompt is designed for single paragraphs, so we use our batch prompt here
            batch_prompt = """You are an expert Document Intelligence Engineer. Analyze the text segments.

DEFINITIONS:
1. Dynamic: Text that changes per client (Names, Dates, Risk Profiles, Amounts).
2. Static: Legal headers, boilerplate, firm branding.

TASK:
Analyze each paragraph and return a JSON response with this exact structure:
{
  "paragraphs": [
    {
      "index": <paragraph_number_from_input>,
      "is_dynamic": true,
      "variables": [
        {"original": "<exact text to replace>", "var": "variable_name_in_snake_case"}
      ]
    }
  ]
}

For non-dynamic paragraphs, set "is_dynamic": false and "variables": [].

PARAGRAPHS TO ANALYZE:
"""

            # Format paragraphs for batch analysis
            paragraphs_text = "\n\n".join(
                f"Paragraph {idx}:\n{text}"
                for idx, text in chunk
            )

            full_prompt = batch_prompt + paragraphs_text

            logger.info(f"Calling LLM for chunk {chunk_idx + 1} with {len(chunk)} paragraphs")

            response = await client.chat.completions.create(
                model=self._model,
                messages=[{"role": "system", "content": full_prompt}],
                response_format={"type": "json_object"},
                temperature=0.1
            )

            content = response.choices[0].message.content
            logger.info(f"LLM response received for chunk {chunk_idx + 1}: {len(content) if content else 0} chars")

            if not content:
                logger.warning(f"Empty LLM response for chunk {chunk_idx + 1}")
                return []

            data = json.loads(content)
            logger.info(f"Parsed JSON for chunk {chunk_idx + 1}: {len(data.get('paragraphs', []))} paragraphs")

            results = []

            for para_data in data.get("paragraphs", []):
                idx = para_data.get("index")
                if para_data.get("is_dynamic", False):
                    for var in para_data.get("variables", []):
                        original = var.get("original", "")
                        var_name = var.get("var", "")
                        if original and var_name:
                            results.append(
                                DetectedVariable(
                                    original_text=original,
                                    suggested_variable_name=var_name,
                                    rationale="LLM Batch Analysis",
                                    paragraph_index=idx
                                )
                            )

            logger.info(f"Chunk {chunk_idx + 1} extracted {len(results)} variables")
            return results

        except ImportError:
            logger.error("openai library not installed. Run: poetry add openai")
            # Fallback to regex
            results = []
            for idx, text in chunk:
                results.extend(self._regex_analysis(text, idx))
            return results
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse LLM JSON response for chunk {chunk_idx}: {e}")
            # Fallback to regex
            results = []
            for idx, text in chunk:
                results.extend(self._regex_analysis(text, idx))
            return results
        except Exception as e:
            logger.error(f"LLM chunk analysis failed for chunk {chunk_idx}: {e}")
            # Fallback to regex
            results = []
            for idx, text in chunk:
                results.extend(self._regex_analysis(text, idx))
            return results

    async def _analyze_paragraph(
        self, text: str, paragraph_index: int
    ) -> list[DetectedVariable]:
        """Analyze a single paragraph for dynamic content.

        Why preserve paragraph_index: Handles duplicate text occurrences.
        E.g., "client_name" may appear in headers and body - we need precise location.

        Args:
            text: The paragraph text to analyze.
            paragraph_index: The paragraph's position in the document.

        Returns:
            List of DetectedVariable objects found in this paragraph.
        """
        if self._use_llm:
            return await self._call_llm(text, paragraph_index)
        else:
            return self._regex_analysis(text, paragraph_index)
# ... inside TemplateAnalyzer class ...

    async def _call_llm(
        self, text: str, paragraph_index: int
    ) -> list[DetectedVariable]:
        """
        Executes the Few-Shot Chain-of-Thought analysis using OpenAI.
        """
        if not self._openai_api_key:
            logger.warning("No OpenAI API key provided. Falling back to regex.")
            return self._regex_analysis(text, paragraph_index)

        try:
            from openai import AsyncOpenAI
            client = AsyncOpenAI(
                api_key=self._openai_api_key,
                base_url=self._base_url,
            )

            # Use custom prompt if provided, otherwise use default
            prompt_template = self._custom_prompt if self._custom_prompt else TEMPLATE_ANALYSIS_PROMPT
            
            # Combine the Few-Shot Prompt with the User Input
            full_prompt = f"{prompt_template}\n\nInput: \"{text}\"\nOutput:"

            response = await client.chat.completions.create(
                model=self._model,
                messages=[{"role": "system", "content": full_prompt}],
                response_format={"type": "json_object"}, # CRITICAL: Enforce JSON
                temperature=0.1 # Low temp for deterministic code/json
            )

            content = response.choices[0].message.content
            if not content:
                return []

            # Parse JSON: { "is_dynamic": bool, "extraction": [...] }
            data = json.loads(content)
            
            if not data.get("is_dynamic", False):
                return []

            results = []
            for item in data.get("extraction", []):
                results.append(
                    DetectedVariable(
                        original_text=item.get("original"),
                        suggested_variable_name=item.get("var"),
                        rationale="LLM Few-Shot Analysis",
                        paragraph_index=paragraph_index
                    )
                )
            return results

        except ImportError:
            logger.error("openai library not installed. Run: poetry add openai")
            return self._regex_analysis(text, paragraph_index)
        except Exception as e:
            logger.error(f"LLM Analysis failed for paragraph {paragraph_index}: {e}")
            # Graceful degradation: If LLM fails, try regex
            return self._regex_analysis(text, paragraph_index)
        
    def _regex_analysis(
        self, text: str, paragraph_index: int
    ) -> list[DetectedVariable]:
        """Analyze text using regex patterns (fallback method).

        Detects common patterns in UK financial documents:
        - Client names (Mr/Mrs/Ms + First + Last)
        - Dates (various formats)
        - Monetary amounts (with £, €, $ symbols)
        - Percentages
        - Address patterns

        Args:
            text: The paragraph text to analyze.
            paragraph_index: The paragraph's position.

        Returns:
            List of DetectedVariable objects detected via regex.
        """
        variables = []

        for pattern, var_name in self._patterns.items():
            try:
                matches = re.findall(pattern, text, re.IGNORECASE)
                for match in matches:
                    # Handle tuple matches from regex groups
                    match_str = (
                        " ".join(match) if isinstance(match, tuple) else str(match)
                    )

                    # Avoid duplicates within same paragraph
                    if any(
                        v.original_text == match_str and v.paragraph_index == paragraph_index
                        for v in variables
                    ):
                        continue

                    variables.append(
                        DetectedVariable(
                            original_text=match_str,
                            suggested_variable_name=var_name,
                            rationale=f"Detected pattern matching {var_name}",
                            paragraph_index=paragraph_index,
                        )
                    )
            except re.error as e:
                logger.warning(f"Regex pattern error for {var_name}: {e}")

        return variables

    @property
    def supported_extensions(self) -> set[str]:
        """Return supported file extensions."""
        return {".docx"}

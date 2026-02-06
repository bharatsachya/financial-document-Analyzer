"""Template injector strategy.

Injects Jinja2 tags or actual values into Word templates at detected
variable locations while preserving original formatting and styles.
"""

import logging
from pathlib import Path
from typing import Any

from app.interfaces.template import BaseTemplateInjector
from app.strategies.template_engine.models import DetectedVariable

logger = logging.getLogger(__name__)


class TemplateInjector(BaseTemplateInjector):
    """Injects Jinja2 tags or actual values into Word templates.

    Uses python-docx to preserve document styles while replacing
    detected variables with Jinja2 template tags or actual values.
    """

    async def inject_tags(
        self,
        file_path: str,
        variables: list[Any],
        output_path: str | None = None,
    ) -> str:
        """Inject Jinja2 tags into the template.

        Why paragraph_index matters: When the same text appears multiple times
        (e.g., "client_name" in header, body, footer), we need to know which
        occurrence to replace. The index ensures precise, safe replacement.

        Args:
            file_path: Path to the original template.
            variables: List of DetectedVariable objects to inject.
            output_path: Where to save the tagged template (optional).

        Returns:
            Path to the generated template file.

        Raises:
            FileNotFoundError: If template file doesn't exist.
            RuntimeError: If injection fails.
        """
        logger.info(f"Starting tag injection: {file_path}")

        try:
            # Validate file exists
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"Template file not found: {file_path}")

            # Load document using python-docx
            try:
                from docx import Document
            except ImportError:
                logger.error("python-docx not installed. Run: poetry add python-docx")
                raise ImportError(
                    "python-docx is required for template injection. "
                    "Install it with: poetry add python-docx"
                )

            doc = Document(file_path)

            # Group variables by paragraph index for efficient processing
            vars_by_paragraph: dict[int, list[Any]] = {}
            for var in variables:
                if var.paragraph_index not in vars_by_paragraph:
                    vars_by_paragraph[var.paragraph_index] = []
                vars_by_paragraph[var.paragraph_index].append(var)

            logger.info(
                f"Processing {len(variables)} variables across "
                f"{len(vars_by_paragraph)} paragraphs"
            )

            # Process each paragraph
            replacement_count = 0
            for idx, paragraph in enumerate(doc.paragraphs):
                if idx not in vars_by_paragraph:
                    continue

                # Apply all replacements for this paragraph
                for var in vars_by_paragraph[idx]:
                    if var.original_text in paragraph.text:
                        # Preserve paragraph runs (styles)
                        self._replace_in_paragraph(paragraph, var)
                        replacement_count += 1
                        logger.debug(
                            f"Replaced at paragraph {idx}: {var.original_text} "
                            f"-> {{% raw %}}{{{{ var.suggested_variable_name }}}}{{% endraw %}}"
                        )

            # Generate output path if not provided
            if output_path is None:
                output_path = str(
                    path.parent / f"{path.stem}_tagged{path.suffix}"
                )

            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)

            # Save the modified document
            doc.save(output_path)
            logger.info(
                f"Tagged template saved: {output_path} ({replacement_count} replacements)"
            )

            return output_path

        except FileNotFoundError:
            raise
        except ImportError:
            raise
        except Exception as e:
            logger.error(f"Tag injection failed: {e}", exc_info=True)
            raise RuntimeError(f"Tag injection failed: {e}") from e

    async def inject_values(
        self,
        file_path: str,
        variables: list[dict[str, Any]],
        output_path: str | None = None,
    ) -> str:
        """Inject actual values into the template (for document generation).

        Replaces detected text with actual user-provided values instead of
        Jinja2 template tags. This is used when generating final documents
        with filled-in data.

        Args:
            file_path: Path to the original template.
            variables: List of dicts with original_text and replacement_value.
            output_path: Where to save the filled document (optional).

        Returns:
            Path to the generated document file.

        Raises:
            FileNotFoundError: If template file doesn't exist.
            RuntimeError: If injection fails.
        """
        logger.info(f"Starting value injection: {file_path}")

        try:
            # Validate file exists
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"Template file not found: {file_path}")

            # Load document using python-docx
            try:
                from docx import Document
            except ImportError:
                logger.error("python-docx not installed. Run: poetry add python-docx")
                raise ImportError(
                    "python-docx is required for template injection. "
                    "Install it with: poetry add python-docx"
                )

            doc = Document(file_path)

            # Create a map of original_text to replacement_value
            value_map: dict[str, str] = {}
            for var in variables:
                original_text = var.get("original_text", "")
                replacement_value = var.get("replacement_value", "")
                if original_text and replacement_value:
                    value_map[original_text] = replacement_value

            logger.info(f"Processing {len(value_map)} value replacements")

            # Process each paragraph
            replacement_count = 0
            for idx, paragraph in enumerate(doc.paragraphs):
                for original_text, replacement_value in value_map.items():
                    if original_text in paragraph.text:
                        self._replace_value_in_paragraph(paragraph, original_text, replacement_value)
                        replacement_count += 1
                        logger.debug(
                            f"Replaced at paragraph {idx}: {original_text} -> {replacement_value}"
                        )

            # Also process tables for values
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for original_text, replacement_value in value_map.items():
                            if original_text in cell.text:
                                for paragraph in cell.paragraphs:
                                    if original_text in paragraph.text:
                                        self._replace_value_in_paragraph(paragraph, original_text, replacement_value)
                                        replacement_count += 1
                                        logger.debug(
                                            f"Replaced in table {table_idx}, row {row_idx}, cell {cell_idx}: "
                                            f"{original_text} -> {replacement_value}"
                                        )

            # Generate output path if not provided
            if output_path is None:
                output_path = str(
                    path.parent / f"{path.stem}_filled{path.suffix}"
                )

            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)

            # Save the modified document
            doc.save(output_path)
            logger.info(
                f"Filled document saved: {output_path} ({replacement_count} replacements)"
            )

            return output_path

        except FileNotFoundError:
            raise
        except ImportError:
            raise
        except Exception as e:
            logger.error(f"Value injection failed: {e}", exc_info=True)
            raise RuntimeError(f"Value injection failed: {e}") from e

    def _replace_in_paragraph(self, paragraph, variable: DetectedVariable) -> None:
        """Replace text in paragraph while preserving styles.

        python-docx uses runs - styled text segments. We must find the run
        containing our text and replace only within it to preserve formatting.

        Args:
            paragraph: The python-docx paragraph object.
            variable: The DetectedVariable to inject.
        """
        # Jinja2 tag for docxtpl (uses {{ var }} syntax)
        jinja_tag = f"{{{{{variable.suggested_variable_name}}}}}"

        # Iterate through runs to find and replace the text
        for run in paragraph.runs:
            if variable.original_text in run.text:
                run.text = run.text.replace(variable.original_text, jinja_tag)
                break  # Only replace in first matching run

    def _replace_value_in_paragraph(
        self, paragraph, original_text: str, replacement_value: str
    ) -> None:
        """Replace text in paragraph with actual value while preserving styles.

        Args:
            paragraph: The python-docx paragraph object.
            original_text: The original text to find and replace.
            replacement_value: The actual value to replace with.
        """
        # Iterate through runs to find and replace the text
        for run in paragraph.runs:
            if original_text in run.text:
                run.text = run.text.replace(original_text, replacement_value)
                break  # Only replace in first matching run

    @property
    def supported_extensions(self) -> set[str]:
        """Return supported file extensions."""
        return {".docx"}
